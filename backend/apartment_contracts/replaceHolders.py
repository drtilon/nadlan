import os
from docx import Document
from docx2pdf import convert


def generate_contract(template_path, output_pdf_path, tenant_details):
    """
    Reads an apartment contract template in DOCX format, replaces placeholders with tenant details,
    and saves the modified document as a new PDF file.

    :param template_path: Path to the contract template DOCX file.
    :param output_pdf_path: Path to save the generated PDF contract.
    :param tenant_details: Dictionary containing tenant-specific details.
    """
    if not os.path.exists(template_path):
        print("Error: Template file does not exist.")
        return

    # Load the DOCX template
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in tenant_details.items():
            placeholder_tag = f"{{{placeholder}}}"
            if placeholder_tag in paragraph.text:
                # Update each run in the paragraph
                for run in paragraph.runs:
                    run.text = run.text.replace(placeholder_tag, value)

    # Replace placeholders in tables (if your template includes tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in tenant_details.items():
                        placeholder_tag = f"{{{placeholder}}}"
                        if placeholder_tag in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(placeholder_tag, value)

    # Save the modified DOCX to a temporary file
    temp_docx = "temp_contract.docx"
    doc.save(temp_docx)

    # Convert the temporary DOCX file to PDF
    try:
        convert(temp_docx, output_pdf_path)
        print(f"Contract generated successfully: {output_pdf_path}")
    except Exception as e:
        print("Error converting DOCX to PDF:", e)
    finally:
        # Clean up the temporary DOCX file
        if os.path.exists(temp_docx):
            os.remove(temp_docx)


# Example usage:
template_file = "apartment_contract_template.docx"
output_pdf_file = "tenant_contract.pdf"
tenant_info = {
    "APARTMENT": "beit lehem 47",
    "LANDLORD_COMPANY_NAME": "Kigel ltd",
    "LANDLORD_NAME": "Kigel ltd",
    "LANDLORD_COMPANY_ADDRESS": "Kigel ltd",
    "LANDLORD_EMAIL": "Kigel ltd",
    "TENANTS_INFO": "John Doe",
    "RENT_AMOUNT": "$1200",
    "LEASE_START": "March 1, 2025",
    "LEASE_END": "February 28, 2026",
}

generate_contract(template_file, output_pdf_file, tenant_info)
