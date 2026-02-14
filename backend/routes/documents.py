import pandas as pd
from io import BytesIO
from datetime import datetime, timedelta
from flask import Blueprint, request, jsonify, g, current_app, send_file, Response
from models.models import Apartment, Tenant, ContractTemplate
from extentions import db
from typing import Tuple, List, Optional
from schemas import ApartmentData, TenantData
from pydantic import ValidationError, BaseModel
from .auth import token_required, role_required
from docx import Document
from docx.shared import Pt, Cm
from num2words import num2words
from utils.logging_helpers import log_with_user
import os

documents_bp = Blueprint("documents_bp", __name__)


# Minimal request model only to get the apartmentId from the client
class ContractRequest(BaseModel):
    apartmentId: int
    templateId: Optional[int] = None  # Optional template ID


@documents_bp.route("/createContract", methods=["POST"])
def create_contract_route() -> tuple[Response, int]:
    try:
        # Get request data (expecting apartmentId and optionally templateId)
        data = request.get_json()
        if not data or "apartmentId" not in data:
            return jsonify({"message": "Invalid request: No apartmentId provided"}), 400

        try:
            # Validate the request; ignore any extra fields
            contract_req = ContractRequest(**data)
        except ValidationError as e:
            return jsonify({"message": "Invalid data", "errors": e.errors()}), 400

        # Retrieve the apartment from the DB
        apartment = Apartment.query.get(contract_req.apartmentId)
        if not apartment:
            return jsonify({"message": "Apartment not found"}), 404

        # Always use tenants associated with the apartment (ignoring client-supplied tenant IDs)
        tenants_list = apartment.tenants
        if not tenants_list:
            return jsonify({"message": "No tenants linked to apartment"}), 400

        # Get the contract template
        template_path = get_contract_template_path(contract_req.templateId)
        if not template_path:
            return jsonify({"message": "Contract template not found"}), 500

        # Use DB dates: fall back to defaults if not set on the apartment
        today, start_date, end_date = format_dates(apartment)
        # Use DB financial details (rent and deposit)
        rent_amount, security_deposit = get_financial_details(apartment)
        # Format tenant details from the DB
        formatted_tenants = format_tenants(tenants_list)

        # Build the dictionary of data to populate the contract template
        db_data = build_db_data(
            apartment,
            rent_amount,
            security_deposit,
            start_date,
            end_date,
            today,
            formatted_tenants,
        )

        # Generate the contract document
        buffer = BytesIO()
        update_rental_agreement(template_path, buffer, db_data)
        buffer.seek(0)

        # Create a filename using apartment address and tenant names
        filename = generate_filename(apartment, tenants_list)

        return send_file(
            buffer,
            download_name=filename,
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        log_with_user(current_app.logger, 'error', f"Error generating contract: {e}")
        return jsonify({"message": "Error generating contract", "error": str(e)}), 500


def get_contract_template_path(template_id=None):
    """
    Get the path to the contract template file.
    If template_id is provided, use that specific template.
    Otherwise, use the default template.
    Falls back to the legacy contract.docx if no templates are found.
    """
    try:
        if template_id:
            # Get specific template by ID
            template = ContractTemplate.query.get(template_id)
            if template and template.file_path and os.path.exists(template.file_path):
                current_app.logger.info(f"Using specific template: {template.name} (ID: {template_id})")
                return template.file_path
            else:
                current_app.logger.warning(f"Requested template ID {template_id} not found or file missing")

        # Get default template
        default_template = ContractTemplate.query.filter_by(is_default=True).first()
        if default_template and default_template.file_path and os.path.exists(default_template.file_path):
            current_app.logger.info(f"Using default template: {default_template.name}")
            return default_template.file_path

        # Get any available template
        any_template = ContractTemplate.query.first()
        if any_template and any_template.file_path and os.path.exists(any_template.file_path):
            current_app.logger.info(f"Using first available template: {any_template.name}")
            return any_template.file_path

        # Fall back to legacy template
        legacy_template_path = os.path.join(current_app.root_path, "contract.docx")
        if os.path.exists(legacy_template_path):
            current_app.logger.warning("Using legacy contract.docx template")
            return legacy_template_path

        # No template found
        log_with_user(current_app.logger, 'error', "No contract template found")
        return None

    except Exception as e:
        log_with_user(current_app.logger, 'error', f"Error getting contract template: {e}")
        # Fall back to legacy template as last resort
        legacy_template_path = os.path.join(current_app.root_path, "contract.docx")
        if os.path.exists(legacy_template_path):
            current_app.logger.warning("Falling back to legacy contract.docx template due to error")
            return legacy_template_path
        return None


def format_dates(apartment):
    """
    Format the start, end, and today's date using values from the apartment.
    Falls back to current date and one year from now if not set.
    """
    today = datetime.now().strftime("%d.%m.%Y")
    if apartment.moveInDate:
        start_date = apartment.moveInDate.strftime("%d.%m.%Y")
    else:
        start_date = today
    if apartment.contractEndDate:
        end_date = apartment.contractEndDate.strftime("%d.%m.%Y")
    else:
        end_date = (datetime.now() + timedelta(days=365)).strftime("%d.%m.%Y")
    return today, start_date, end_date


def get_financial_details(apartment):
    """
    Retrieve the rent amount and security deposit from the apartment record.
    """
    rent_amount = apartment.rent
    security_deposit = apartment.deposit
    return rent_amount, security_deposit


def format_tenants(tenants_list):
    """
    Format tenant details from the DB into a list of dictionaries suitable for
    insertion into the contract template.
    """
    formatted = []
    for tenant in tenants_list:
        dob = ""
        if tenant.bornOn:
            try:
                if isinstance(tenant.bornOn, str):
                    dob_date = datetime.fromisoformat(tenant.bornOn)
                else:
                    dob_date = tenant.bornOn
                dob = dob_date.strftime("%d.%m.%Y")
            except Exception as e:
                dob = tenant.bornOn  # fallback if parsing fails
        formatted.append(
            {
                "name": tenant.name,
                "dob": dob,
                "phone": tenant.phone or "",
                "email": tenant.email or "",
            }
        )
    return formatted


def build_db_data(
    apartment,
    rent_amount,
    security_deposit,
    start_date,
    end_date,
    today,
    formatted_tenants,
):
    """
    Build the dictionary of data to pass to the contract template updater,
    using all values from the database.
    """

    return {
        "landlord_company_name": apartment.landlord.name if apartment.landlord else "",
        "landlord_name": apartment.landlord.company_name,
        "landlord_company_address": apartment.landlord.company_address
        if apartment.landlord
        else "",
        "landlord_email": apartment.landlord.email if apartment.landlord else "",
        "landlord_phone": apartment.landlord.phone if apartment.landlord else "",
        "landlord_iban": apartment.landlord.iban if apartment.landlord else "",
        "apartment_address": apartment.address,
        "rent_price": f"{apartment.rent:.2f}",
        "rent_words": str(num2words(int(apartment.rent))),
        "deposit": apartment.deposit,
        "start_date": start_date,
        "end_date": end_date,
        "today_date": today,
        "tenants": formatted_tenants,
        "special_terms": apartment.notes or "",
    }


def generate_filename(apartment, tenants_list):
    """
    Generate a filename for the contract document based on the apartment address
    and tenant names.
    """
    if len(tenants_list) == 1:
        tenant_part = tenants_list[0].name.replace(" ", "_")
    else:
        tenant_part = (
            f"{tenants_list[0].name.replace(' ', '_')}_and_{len(tenants_list) - 1}_more"
        )
    return f"Rental_Contract_{apartment.address.replace(' ', '_')}_{tenant_part}.docx"


def update_rental_agreement(template_path, output, db_data):
    """
    Updates a rental agreement template with data from the database.
    """
    doc = Document(template_path)

    # Define placeholders mapped to the data from the DB entities
    placeholders = {
        "{LANDLORD_COMPANY_NAME}": db_data["landlord_company_name"],
        "{LANDLORD_NAME}": db_data["landlord_name"],
        "{LANDLORD_COMPANY_ADDRESS}": db_data["landlord_company_address"],
        "{LANDLORD_EMAIL}": db_data["landlord_email"],
        "{LANDLORD_PHONE}": db_data["landlord_phone"],
        "{LANDLORD_IBAN}": db_data["landlord_iban"],
        "{APARTMENT_ADDRESS}": db_data["apartment_address"],
        "{RENT_PRICE}": db_data["rent_price"],
        "{RENT_PRICE_WORDS}": db_data["rent_words"],
        "{DEPOSIT}": db_data["deposit"],
        "{START_DATE}": db_data["start_date"],
        "{END_DATE}": db_data["end_date"],
        "{TODAY_DATE}": db_data["today_date"],
        "{SPECIAL_TERMS}": db_data["special_terms"],
    }

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in para.text:
                replace_text_in_paragraph(para, placeholder, str(value))

    # Replace placeholders in all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        if placeholder in para.text:
                            replace_text_in_paragraph(para, placeholder, str(value))

    # Look for the tenant placeholder "{TENANT_LIST}" in document paragraphs or tables
    tenant_placeholder = "{TENANT_LIST}"
    tenant_para = None

    # First, try in document paragraphs
    for para in doc.paragraphs:
        if tenant_placeholder in para.text:
            tenant_para = para
            break

    # If not found, search in table cells
    if not tenant_para:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if tenant_placeholder in para.text:
                            tenant_para = para
                            break
                    if tenant_para:
                        break
                if tenant_para:
                    break
            if tenant_para:
                break

    # If a tenant placeholder is found, remove it and insert tenant details
    if tenant_para:
        parent = tenant_para._element.getparent()
        index = parent.index(tenant_para._element)
        style_name = tenant_para.style.name
        parent.remove(tenant_para._element)

        for idx, tenant in enumerate(db_data["tenants"], start=1):
            p = doc.add_paragraph(style=style_name)
            tenant_text = f"{idx}. {tenant['name']}"
            details = []
            if tenant["dob"]:
                details.append(f"born on {tenant['dob']}")
            if tenant["phone"]:
                details.append(f"Phone: {tenant['phone']}")
            if tenant["email"]:
                details.append(f"E-mail: {tenant['email']}")
            if details:
                tenant_text += ", " + ", ".join(details)
            p.add_run(tenant_text)
            parent.insert(index, p._element)
            index += 1

    doc.save(output)


def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """
    Replace all occurrences of a placeholder in a paragraph, even if the
    placeholder spans across multiple runs.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return
    new_text = full_text.replace(placeholder, replacement)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
